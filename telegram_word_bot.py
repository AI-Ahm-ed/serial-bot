import os
import logging
from telegram import Update
from telegram.ext import ApplicationBuilder, ContextTypes, MessageHandler, filters
from docx import Document

# إعداد السجلات للتتبع
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# التوكن الخاص بك
TOKEN = "8876238881:AAEVbcBHKdpsFRIHxj_P5me6NLEc0JXA2lU"

# قائمة مؤقتة لتخزين الكارتات المرسلة
cards_database = []

def copy_paragraph(source_p, target_doc):
    """نسخ فقرة بكل تنسيقاتها إلى المستند الهدف"""
    new_p = target_doc.add_paragraph()
    new_p.alignment = source_p.alignment
    new_p.style = source_p.style
    for run in source_p.runs:
        new_run = new_p.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        new_run.font.color.rgb = run.font.color.rgb
    return new_p

def copy_table(source_table, target_doc):
    """نسخ جدول بكل محتوياته إلى المستند الهدف"""
    new_table = target_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    new_table.style = source_table.style
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.cell(i, j)
            new_cell.text = cell.text
    return new_table

def create_combined_word_document(cards_list):
    """
    دمج قالب الوورد في ملف واحد بحيث يأخذ كل كارت صفحة جديدة مستقلة تماماً ومنظمة
    """
    base_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(base_dir, "template.docx")
    
    if not os.path.exists(template_path):
        logger.error(f"ملف القالب غير موجود في المسار: {template_path}")
        return None

    try:
        template_doc = Document(template_path)
    except Exception as e:
        logger.error(f"فشل في فتح ملف القالب: {e}")
        return None

    # إنشاء مستند جديد خالي تماماً لنبني فيه الكارتات نظيفة
    master_doc = Document()

    for index, card in enumerate(cards_list):
        category = card['category']
        serial = card['serial']
        pin = card['pin']
        exp = card['exp']

        # إذا لم يكن الكارت الأول، نضيف فاصل صفحات حقيقي ونظيف
        if index > 0:
            master_doc.add_page_break()

        # قراءة عناصر القالب وترتيبها كفقرات وجداول داخل الصفحة الحالية
        for element in template_doc.element.body:
            # التحقق مما إذا كان العنصر فقرة
            if element.tag.endswith('p'):
                # إنشاء فقرة مؤقتة لمعالجة النصوص واستبدال المتغيرات
                temp_p_text = "".join([node.text for node in element.iter() if node.text])
                
                # استبدال الحقول مباشرة
                filled_text = (temp_p_text
                               .replace("[CATEGORY]", category)
                               .replace("[SERIAL]", serial)
                               .replace("[PIN]", pin)
                               .replace("[EXP]", exp))
                
                master_doc.add_paragraph(filled_text)
                
            # التحقق مما إذا كان العنصر جدولاً
            elif element.tag.endswith('tbl'):
                # البحث عن الجدول المطابق في القالب ونسخه مع استبدال البيانات بداخله
                for table in template_doc.tables:
                    if table._tbl == element:
                        new_table = master_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table.style = table.style
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                cell_text = cell.text
                                filled_cell_text = (cell_text
                                                    .replace("[CATEGORY]", category)
                                                    .replace("[SERIAL]", serial)
                                                    .replace("[PIN]", pin)
                                                    .replace("[EXP]", exp))
                                new_table.cell(i, j).text = filled_cell_text
                        break

    output_filename = "all_cards_combined.docx"
    output_path = os.path.join(base_dir, output_filename)
    master_doc.save(output_path)
    return output_path

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    global cards_database
    text = update.message.text.strip()
    
    # 1. أمر تفريغ القائمة
    if text.lower() == "clear":
        count = len(cards_database)
        cards_database = []
        await update.message.reply_text(f"🗑️ تم مسح وتفريغ القائمة بنجاح! (تم حذف {count} كارت من الذاكرة).")
        return

    # 2. أمر طلب الملف المجمع
    if text.lower() == "file":
        if not cards_database:
            await update.message.reply_text("⚠️ لم تقم بإرسال أي كارت بعد! أرسل الكارتات أولاً ثم اكتب file.")
            return
        
        await update.message.reply_text(f"⏳ جاري دمج {len(cards_database)} كارت في ملف واحد...")
        
        file_path = create_combined_word_document(cards_database)
        
        if file_path and os.path.exists(file_path):
            with open(file_path, 'rb') as doc_file:
                await update.message.reply_document(
                    document=doc_file,
                    caption=f"✅ تم إرسال الملف المجمع ويحتوي على ({len(cards_database)}) كارت في صفحات مستقلة تماماً.\n(ملاحظة: الكارتات لا تزال محفوظة، يمكنك طلب الملف مرة أخرى أو كتابة clear للتفريغ)."
                )
            os.remove(file_path)
        else:
            await update.message.reply_text("❌ حدث خطأ أثناء دمج المستندات، تأكد من وجود ملف `template.docx`.")
        return

    # 3. معالجة وتخزين الكارت الوارد
    lines = text.split('\n')
    
    if len(lines) < 4:
        await update.message.reply_text(
            "⚠️ الصيغة غير صحيحة!\n"
            "الرجاء إرسال الرسالة بنفس التنسيق:\n\n"
            "15K\n"
            "Ser: 26041900225094\n"
            "PIN: 2421896018520315\n"
            "Exp: 2028-08-31\n\n"
            "الأوامر المتاحة:\n"
            "- **file**: لاستلام الملف المجمع.\n"
            "- **clear**: لتفريغ القائمة وحذف الكارتات المخزنة."
        )
        return

    category = lines[0].strip()
    serial = ""
    pin = ""
    exp = ""

    for line in lines[1:]:
        if "Ser:" in line:
            serial = line.replace("Ser:", "").strip()
        elif "PIN:" in line:
            pin = line.replace("PIN:", "").strip()
        elif "Exp:" in line:
            exp = line.replace("Exp:", "").strip()

    if not serial or not pin or not exp:
        await update.message.reply_text("⚠️ لم يتم العثور على جميع البيانات (Ser, PIN, Exp). تأكد من كتابة الكلمات بشكل صحيح.")
        return

    cards_database.append({
        'category': category,
        'serial': serial,
        'pin': pin,
        'exp': exp
    })

    await update.message.reply_text(f"📥 تم حفظ الكارت (Ser: {serial})\nالعدد الإجمالي المخزن حالياً: {len(cards_database)}")

def main():
    application = ApplicationBuilder().token(TOKEN).build()
    application.add_handler(MessageHandler(filters.TEXT & (~filters.COMMAND), handle_message))

    print("Bot is running...")
    application.run_polling()

if __name__ == '__main__':
    main()
